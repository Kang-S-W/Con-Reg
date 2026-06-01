import streamlit as st
import streamlit.components.v1 as st_components
import io
import textwrap
import html
import os
import uuid
import pandas as pd
from docx import Document

# ==========================================
# 1. 페이지 설정
# ==========================================
st.set_page_config(
    page_title="용인시 건축 조례 지원 플랫폼",
    layout="wide",
    initial_sidebar_state="expanded",
)

# 외부 모듈 임포트
from widget_utils import inject_floating_button
from processor import handle_ai_analysis, generate_civil_document
from style import apply_custom_style
from components import render_user_message, render_ai_report
from storage import load_history, save_history

# ==========================================
# 2. 세션 상태 관리 및 초기화
# ==========================================
if "user_id" not in st.session_state:
    st.session_state.user_id = "guest"
if "chat_history" not in st.session_state:
    st.session_state.chat_history = load_history(st.session_state.user_id)
if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = False
if "dark_mode_toggle" not in st.session_state:
    st.session_state.dark_mode_toggle = st.session_state.dark_mode
if "selected_index" not in st.session_state:
    st.session_state.selected_index = None
if "current_page" not in st.session_state:
    st.session_state.current_page = "main"
if "qna_list" not in st.session_state:
    st.session_state.qna_list = []

if (
    st.session_state.selected_index is not None
    and st.session_state.selected_index >= len(st.session_state.chat_history)
):
    st.session_state.selected_index = None

# 챗봇 메인 화면에서 실제 대화가 있을 때만 민원창구 플로팅 버튼 표시
selected_index = st.session_state.get("selected_index")
chat_history = st.session_state.get("chat_history", [])

if (
    st.session_state.get("current_page") == "main"
    and selected_index is not None
    and 0 <= selected_index < len(chat_history)
    and chat_history[selected_index].get("messages")
):
    inject_floating_button()

def sync_dark_mode():
    st.session_state.dark_mode = st.session_state.dark_mode_toggle


apply_custom_style(st.session_state.dark_mode)


def find_yongin_logo():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    logo_candidates = [
        "yongin_logo.png",
        "Con-Reg/yongin_logo.png",
        "assets/yongin_logo.png",
        "assets/yongin_logo.jpg",
        "assets/yongin_logo.jpeg",
        "assets/yongin_logo.svg",
        "assets/yongin-city-logo.png",
        "assets/yongin-city-logo.svg",
        "assets/logo.png",
        "assets/logo.jpg",
        "assets/logo.svg",
        "yongin_logo.jpg",
        "yongin_logo.svg",
        "yongin-city-logo.png",
        "yongin-city-logo.svg",
        "용인시_로고.png",
        "용인시_로고.jpg",
        "용인시_로고.svg",
        "logo.png",
        "logo.jpg",
        "logo.svg",
    ]
    for candidate in logo_candidates:
        logo_path = os.path.join(base_dir, candidate)
        if os.path.exists(logo_path):
            return logo_path

    base_depth = base_dir.rstrip(os.sep).count(os.sep)
    for root, dirs, files in os.walk(base_dir):
        if root.rstrip(os.sep).count(os.sep) - base_depth >= 3:
            dirs[:] = []
        if "yongin_logo.png" in files:
            return os.path.join(root, "yongin_logo.png")
    return None


def ensure_chat_ids():
    if not st.session_state.chat_history:
        return

    changed = False
    for chat in st.session_state.chat_history:
        if "id" not in chat:
            chat["id"] = str(uuid.uuid4())
            changed = True

    if changed:
        save_history(st.session_state.chat_history, st.session_state.user_id)


# ==========================================
# 3. 전문 업무용 UI 테마
# ==========================================
def apply_premium_ui_v3(is_dark):
    if is_dark:
        colors = {
            "app_bg": "#121212",
            "sidebar_bg": "#151515",
            "surface": "#1f1f1f",
            "surface_2": "#282828",
            "surface_3": "#333333",
            "text": "#f5f5f5",
            "muted": "#b8b8b8",
            "line": "#4b4b4b",
            "accent": "#e7e7e7",
            "accent_strong": "#ffffff",
            "accent_soft": "rgba(255, 255, 255, 0.08)",
            "success_bg": "rgba(255, 255, 255, 0.06)",
            "success_line": "#bdbdbd",
            "warning_bg": "rgba(255, 255, 255, 0.06)",
            "warning_line": "#bdbdbd",
            "error_bg": "rgba(255, 255, 255, 0.06)",
            "error_line": "#bdbdbd",
            "button_bg": "#f7f7f7",
            "button_text": "#111111",
            "button_hover_bg": "#ffffff",
            "table_bg": "#1f1f1f",
            "table_header": "#282828",
            "table_text": "#f5f5f5",
            "table_muted": "#b8b8b8",
            "toggle_off": "#ef4444",
            "toggle_on": "#22c55e",
            "toggle_thumb": "#ffffff",
            "toggle_border": "#f5f5f5",
        }
    else:
        colors = {
            "app_bg": "#f5f5f5",
            "sidebar_bg": "#f1f1f1",
            "surface": "#ffffff",
            "surface_2": "#fafafa",
            "surface_3": "#eeeeee",
            "text": "#111111",
            "muted": "#5b5b5b",
            "line": "#b8b8b8",
            "accent": "#1f1f1f",
            "accent_strong": "#000000",
            "accent_soft": "#eeeeee",
            "success_bg": "#f7f7f7",
            "success_line": "#9a9a9a",
            "warning_bg": "#f7f7f7",
            "warning_line": "#9a9a9a",
            "error_bg": "#f7f7f7",
            "error_line": "#9a9a9a",
            "button_bg": "#ffffff",
            "button_text": "#111111",
            "button_hover_bg": "#f4f4f4",
            "table_bg": "#ffffff",
            "table_header": "#f2f2f2",
            "table_text": "#111111",
            "table_muted": "#555555",
            "toggle_off": "#ef4444",
            "toggle_on": "#22c55e",
            "toggle_thumb": "#ffffff",
            "toggle_border": "#111111",
        }

    st.markdown(
        f"""
<style>
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');

    :root {{
        --app-bg: {colors["app_bg"]};
        --sidebar-bg: {colors["sidebar_bg"]};
        --surface: {colors["surface"]};
        --surface-2: {colors["surface_2"]};
        --surface-3: {colors["surface_3"]};
        --text: {colors["text"]};
        --muted: {colors["muted"]};
        --line: {colors["line"]};
        --accent: {colors["accent"]};
        --accent-strong: {colors["accent_strong"]};
        --accent-soft: {colors["accent_soft"]};
        --success-bg: {colors["success_bg"]};
        --success-line: {colors["success_line"]};
        --warning-bg: {colors["warning_bg"]};
        --warning-line: {colors["warning_line"]};
        --error-bg: {colors["error_bg"]};
        --error-line: {colors["error_line"]};
        --button-bg: {colors["button_bg"]};
        --button-text: {colors["button_text"]};
        --button-hover-bg: {colors["button_hover_bg"]};
        --table-bg: {colors["table_bg"]};
        --table-header: {colors["table_header"]};
        --table-text: {colors["table_text"]};
        --table-muted: {colors["table_muted"]};
        --toggle-off: {colors["toggle_off"]};
        --toggle-on: {colors["toggle_on"]};
        --toggle-thumb: {colors["toggle_thumb"]};
        --toggle-border: {colors["toggle_border"]};
    }}

    html, body, .stApp, p, h1, h2, h3, h4, h5, h6, label, input, textarea, div {{
        font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, sans-serif;
        letter-spacing: 0;
    }}

    .stApp,
    [data-testid="stAppViewContainer"] {{
        background: var(--app-bg) !important;
        color: var(--text) !important;
    }}

    .block-container {{
        max-width: 1320px;
        padding-top: 2rem;
        padding-bottom: 3rem;
    }}

    h1, h2, h3, h4, h5, h6,
    .stMarkdown, .stMarkdown p, .stCaptionContainer, label, .stText {{
        color: var(--text) !important;
    }}

    .stCaptionContainer,
    small,
    div[data-testid="stMarkdownContainer"] p {{
        color: var(--muted) !important;
        line-height: 1.65 !important;
    }}

    [data-testid="stSidebar"],
    [data-testid="stSidebar"] > div {{
        background: var(--sidebar-bg) !important;
        border-right: 1.25px solid var(--line) !important;
    }}

    [data-testid="stSidebar"] > div:first-child {{
        padding-top: 0 !important;
    }}

    [data-testid="stSidebarUserContent"] {{
        padding-top: 0 !important;
    }}

    [data-testid="stSidebar"] [data-testid="stVerticalBlock"] {{
        gap: 0.45rem !important;
    }}

    .sidebar-brand {{
        text-align: center;
        padding: 0 4px 8px;
    }}

    .sidebar-brand-title {{
        margin-top: -2px;
        color: var(--accent-strong);
        font-size: 20px;
        font-weight: 850;
        line-height: 1.32;
        text-align: center;
    }}

    [data-testid="stSidebar"] [data-testid="stImage"] {{
        text-align: center;
    }}

    [data-testid="stSidebar"] [data-testid="stImage"] img {{
        display: block;
        margin: 0 auto;
        max-height: 82px;
        object-fit: contain;
        image-rendering: auto;
        transform: translate(42px, -8px);
    }}

    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3,
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] label {{
        color: var(--text) !important;
    }}

    [data-testid="stSidebar"] .stButton > button {{
        min-height: 40px;
        border-radius: 10px;
        font-weight: 650;
    }}

    .stButton > button,
    .stDownloadButton > button,
    button[data-testid="baseButton-secondary"],
    button[data-testid="baseButton-primary"],
    button[kind="primary"],
    button[kind="secondary"],
    button[kind="tertiary"] {{
        background: var(--button-bg) !important;
        color: var(--button-text) !important;
        -webkit-text-fill-color: var(--button-text) !important;
        border: 1.25px solid var(--line) !important;
        border-radius: 10px !important;
        box-shadow: 0 1px 2px rgba(15, 23, 42, 0.06) !important;
        transition: border-color 0.18s ease, background 0.18s ease, transform 0.18s ease !important;
    }}

    div[data-testid="stFormSubmitButton"] button,
    button[kind="secondaryFormSubmit"],
    button[kind="primaryFormSubmit"] {{
        background: var(--button-bg) !important;
        color: var(--button-text) !important;
        -webkit-text-fill-color: var(--button-text) !important;
        border: 1.4px solid var(--line) !important;
        border-radius: 10px !important;
        font-weight: 750 !important;
    }}

    .stButton > button:disabled,
    .stDownloadButton > button:disabled,
    div[data-testid="stFormSubmitButton"] button:disabled,
    button:disabled,
    button[disabled] {{
        background: var(--button-bg) !important;
        color: var(--button-text) !important;
        -webkit-text-fill-color: var(--button-text) !important;
        border-color: var(--line) !important;
        opacity: 1 !important;
        cursor: not-allowed !important;
    }}

    div[data-testid="stFormSubmitButton"] button:hover,
    button[kind="secondaryFormSubmit"]:hover,
    button[kind="primaryFormSubmit"]:hover {{
        background: var(--button-hover-bg) !important;
        color: var(--button-text) !important;
        -webkit-text-fill-color: var(--button-text) !important;
        border-color: var(--accent-strong) !important;
    }}

    .stButton > button:hover,
    .stDownloadButton > button:hover {{
        background: var(--button-hover-bg) !important;
        border-color: var(--accent) !important;
        color: var(--button-text) !important;
        -webkit-text-fill-color: var(--button-text) !important;
        transform: translateY(-1px);
    }}

    button[kind="primary"],
    button[data-testid="baseButton-primary"] {{
        font-weight: 750 !important;
    }}

    div[data-baseweb="input"] > div,
    div[data-baseweb="textarea"] > div,
    div[data-baseweb="select"] > div,
    [data-testid="stChatInput"] > div {{
        background: var(--surface) !important;
        border: 1.6px solid var(--line) !important;
        border-radius: 12px !important;
        color: var(--text) !important;
        box-shadow: none !important;
    }}

    div[data-baseweb="input"] > div:focus-within,
    div[data-baseweb="textarea"] > div:focus-within,
    div[data-baseweb="select"] > div:focus-within,
    [data-testid="stChatInput"] > div:focus-within {{
        border-color: var(--accent) !important;
        box-shadow: 0 0 0 3px var(--accent-soft) !important;
    }}

    input,
    textarea,
    div[data-baseweb="input"] input,
    div[data-baseweb="textarea"] textarea,
    [data-testid="stChatInput"] textarea {{
        color: var(--text) !important;
        -webkit-text-fill-color: var(--text) !important;
        caret-color: var(--accent) !important;
        font-size: 15px !important;
    }}

    input::placeholder,
    textarea::placeholder,
    [data-testid="stChatInput"] textarea::placeholder {{
        color: var(--muted) !important;
        opacity: 0.85 !important;
    }}

    .toggle-caption {{
        color: var(--text) !important;
        -webkit-text-fill-color: var(--text) !important;
        background: transparent !important;
        border: 0 !important;
        outline: 0 !important;
        box-shadow: none !important;
        font-size: 15px;
        font-weight: 750;
        line-height: 1.2;
        padding-top: 2px;
        white-space: nowrap;
    }}

    div[data-testid="stToggle"] label,
    div[data-testid="stCheckbox"] label,
    div[data-testid="stToggle"] [data-testid="stWidgetLabel"],
    div[data-testid="stCheckbox"] [data-testid="stWidgetLabel"],
    div[data-testid="stToggle"] [data-testid="stMarkdownContainer"],
    div[data-testid="stCheckbox"] [data-testid="stMarkdownContainer"],
    div[data-testid="stToggle"] label p,
    div[data-testid="stCheckbox"] label p,
    div[data-testid="stToggle"] label span,
    div[data-testid="stCheckbox"] label span,
    div[data-testid="stToggle"] [data-testid="stWidgetLabel"] *,
    div[data-testid="stCheckbox"] [data-testid="stWidgetLabel"] *,
    div[data-testid="stToggle"] [data-testid="stMarkdownContainer"] *,
    div[data-testid="stCheckbox"] [data-testid="stMarkdownContainer"] * {{
        color: var(--text) !important;
        -webkit-text-fill-color: var(--text) !important;
        background: transparent !important;
        background-color: transparent !important;
        border: 0 !important;
        outline: 0 !important;
        box-shadow: none !important;
        border-radius: 0 !important;
        text-shadow: none !important;
        opacity: 1 !important;
        font-weight: 700 !important;
    }}

    div[data-testid="stToggle"] > label,
    div[data-testid="stCheckbox"] > label {{
        background: transparent !important;
        background-color: transparent !important;
        border: 0 !important;
        outline: 0 !important;
        box-shadow: none !important;
        gap: 0.55rem !important;
    }}

    div[data-testid="stToggle"] [data-testid="stCheckboxIndicator"],
    div[data-testid="stCheckbox"] [data-testid="stCheckboxIndicator"] {{
        background: var(--toggle-off) !important;
        border: 2px solid var(--toggle-border) !important;
        box-shadow: 0 1px 5px rgba(0, 0, 0, 0.22) !important;
        opacity: 1 !important;
    }}

    div[data-testid="stToggle"] [role="switch"],
    div[data-testid="stCheckbox"] [role="switch"] {{
        background: var(--toggle-off) !important;
        border: 2px solid var(--toggle-border) !important;
        box-shadow: 0 1px 5px rgba(0, 0, 0, 0.22) !important;
        opacity: 1 !important;
    }}

    div[data-testid="stToggle"] [role="switch"][aria-checked="true"],
    div[data-testid="stCheckbox"] [role="switch"][aria-checked="true"] {{
        background: var(--toggle-on) !important;
        border-color: var(--toggle-border) !important;
    }}

    div[data-testid="stToggle"] [data-testid="stCheckboxIndicator"]::before,
    div[data-testid="stCheckbox"] [data-testid="stCheckboxIndicator"]::before {{
        background: var(--toggle-thumb) !important;
        opacity: 1 !important;
    }}

    div[data-testid="stToggle"] input:checked + div,
    div[data-testid="stCheckbox"] input:checked + div {{
        background: var(--toggle-on) !important;
        border-color: var(--toggle-border) !important;
    }}

    div[data-testid="stToggle"] input:checked ~ div [data-testid="stCheckboxIndicator"],
    div[data-testid="stCheckbox"] input:checked ~ div [data-testid="stCheckboxIndicator"],
    div[data-testid="stToggle"] input:checked ~ [data-testid="stCheckboxIndicator"],
    div[data-testid="stCheckbox"] input:checked ~ [data-testid="stCheckboxIndicator"] {{
        background: var(--toggle-on) !important;
        border-color: var(--toggle-border) !important;
    }}

    [data-testid="stSidebar"] input[type="checkbox"] + div,
    [data-testid="stSidebar"] input[role="switch"] + div {{
        background: var(--toggle-off) !important;
        border: 2px solid var(--toggle-border) !important;
        box-shadow: 0 1px 5px rgba(0, 0, 0, 0.22) !important;
        opacity: 1 !important;
    }}

    [data-testid="stSidebar"] input[type="checkbox"]:checked + div,
    [data-testid="stSidebar"] input[role="switch"]:checked + div {{
        background: var(--toggle-on) !important;
        border-color: var(--toggle-border) !important;
    }}

    div[data-testid="stExpander"] details {{
        background: var(--surface) !important;
        border: 1.25px solid var(--line) !important;
        border-radius: 12px !important;
        overflow: hidden !important;
    }}

    div[data-testid="stExpander"] summary {{
        background: var(--surface-2) !important;
        color: var(--text) !important;
        font-weight: 700 !important;
        border-bottom: 1.25px solid var(--line) !important;
    }}

    div[data-testid="stExpander"] details > div {{
        background: var(--surface) !important;
        color: var(--text) !important;
    }}

    [data-testid="stVerticalBlockBorderWrapper"] {{
        background: var(--surface) !important;
        border: 1.25px solid var(--line) !important;
        border-radius: 12px !important;
        box-shadow: 0 8px 24px rgba(15, 23, 42, 0.05) !important;
    }}

    div[data-testid="stTabBar"] {{
        gap: 8px !important;
        border-bottom: 1.25px solid var(--line) !important;
        padding-bottom: 8px !important;
    }}

    button[data-baseweb="tab"] {{
        background: var(--surface) !important;
        color: var(--muted) !important;
        border: 1.25px solid var(--line) !important;
        border-radius: 10px !important;
        padding: 8px 14px !important;
        font-weight: 650 !important;
    }}

    button[data-baseweb="tab"][aria-selected="true"] {{
        background: var(--accent-soft) !important;
        color: var(--accent-strong) !important;
        border-color: var(--accent) !important;
    }}

    [data-testid="stDataFrame"],
    [data-testid="stDataEditor"] {{
        border: 1.25px solid var(--line) !important;
        border-radius: 12px !important;
        overflow: hidden !important;
        background: var(--surface) !important;
    }}

    .legal-table-wrap {{
        overflow-x: auto;
        border: 1.35px solid var(--line);
        border-radius: 12px;
        background: var(--table-bg);
        margin-top: 10px;
        margin-bottom: 18px;
    }}

    .legal-table {{
        width: 100%;
        border-collapse: collapse;
        background: var(--table-bg);
        color: var(--table-text);
        font-size: 15px;
    }}

    .legal-table thead th {{
        background: var(--table-header);
        color: var(--table-text);
        font-weight: 800;
        text-align: left;
        border-bottom: 1.25px solid var(--line);
        border-right: 1px solid var(--line);
        padding: 14px 16px;
        white-space: nowrap;
    }}

    .legal-table tbody td {{
        background: var(--table-bg);
        color: var(--table-text);
        border-top: 1px solid var(--line);
        border-right: 1px solid var(--line);
        padding: 13px 16px;
        vertical-align: top;
        line-height: 1.45;
    }}

    .legal-table th:last-child,
    .legal-table td:last-child {{
        border-right: none;
    }}

    [data-testid="stDialog"],
    div[role="dialog"] {{
        background: var(--surface) !important;
        color: var(--text) !important;
        border: 1.25px solid var(--line) !important;
        border-radius: 16px !important;
    }}

    [data-testid="stDialog"] *,
    div[role="dialog"] * {{
        color: var(--text) !important;
    }}

    [data-testid="stDialog"] input,
    [data-testid="stDialog"] textarea,
    div[role="dialog"] input,
    div[role="dialog"] textarea {{
        color: var(--text) !important;
        -webkit-text-fill-color: var(--text) !important;
    }}

    .app-intro {{
        text-align: center;
        padding: 56px 24px;
        border: 1.25px solid var(--line);
        border-radius: 16px;
        background: var(--surface);
    }}

    .app-intro h3 {{
        color: var(--accent-strong) !important;
        font-weight: 750;
        margin-bottom: 12px;
    }}

    .app-intro p {{
        color: var(--muted) !important;
        font-size: 14px;
    }}

    .app-notice {{
        border: 1.25px solid var(--line);
        border-left-width: 4px;
        border-radius: 12px;
        padding: 14px 16px;
        margin: 8px 0 12px;
        color: var(--text);
        background: var(--surface);
        line-height: 1.6;
        font-size: 14px;
    }}

    .app-notice-info {{
        border-left-color: var(--accent);
        background: var(--accent-soft);
    }}

    .app-notice-success {{
        border-left-color: var(--success-line);
        background: var(--success-bg);
    }}

    .app-notice-warning {{
        border-left-color: var(--warning-line);
        background: var(--warning-bg);
    }}

    .app-notice-error {{
        border-left-color: var(--error-line);
        background: var(--error-bg);
    }}

    .app-panel {{
        padding: 18px;
        border: 1.25px solid var(--line);
        border-radius: 14px;
        background: var(--surface);
    }}

    .context-panel-title {{
        font-size: 16px;
        font-weight: 750;
        margin-bottom: 2px;
        color: var(--text);
    }}

    .metric-label {{
        font-size: 13px;
        color: var(--muted);
        margin-bottom: 4px;
    }}

    .metric-title {{
        font-size: 18px;
        font-weight: 750;
        color: var(--text);
        margin-bottom: 6px;
    }}
</style>
""",
        unsafe_allow_html=True,
    )


apply_premium_ui_v3(st.session_state.dark_mode)


def render_notice(message, tone="info"):
    safe_message = html.escape(str(message)).replace("\n", "<br>")
    st.markdown(
        f"<div class='app-notice app-notice-{tone}'>{safe_message}</div>",
        unsafe_allow_html=True,
    )


def to_multi_column_dataframe(items, column_count=4, column_name="법규"):
    rows = []
    for start in range(0, len(items), column_count):
        row = items[start : start + column_count]
        row += [""] * (column_count - len(row))
        rows.append(row)
    return pd.DataFrame(rows, columns=[f"{column_name} {i + 1}" for i in range(column_count)])


def render_legal_table(dataframe):
    table_html = dataframe.fillna("").to_html(
        index=False,
        escape=True,
        classes="legal-table",
        border=0,
    )
    st.markdown(f"<div class='legal-table-wrap'>{table_html}</div>", unsafe_allow_html=True)


# ==========================================
# 4. 대화 기록 검색 팝업
# ==========================================
dialog_decorator = st.dialog if hasattr(st, "dialog") else st.experimental_dialog


@dialog_decorator("대화 기록 검색", width="large")
def open_history_search_dialog():
    search_query = st.text_input(
        "검색어",
        placeholder="예: 건폐율, 사선제한, 처인구",
        key="dialog_history_search_input",
    )
    query = search_query.strip().lower()

    if not st.session_state.chat_history:
        render_notice("저장된 대화가 없습니다.", "info")
        return

    results = []
    for idx, chat in enumerate(st.session_state.chat_history):
        searchable_text = chat.get("title", "") + " "
        for msg in chat.get("messages", []):
            searchable_text += msg.get("query", "") + " " + msg.get("response", "") + " "
        if not query or query in searchable_text.lower():
            results.append((idx, chat))

    st.caption(f"검색 결과: {len(results)}건" if query else "최근 대화")

    if not results:
        render_notice("일치하는 검색 결과가 없습니다.", "warning")
        return

    with st.container(height=420, border=True):
        for idx, chat in reversed(results[-20:]):
            title = chat.get("title", "새 대화")
            time_str = chat.get("updated_at", chat.get("created_at", ""))
            preview = (
                chat.get("messages", [])[0].get("query", "")[:60] + "..."
                if chat.get("messages")
                else ""
            )

            with st.container():
                if st.button(
                    f"{title[:28]}\n\n{time_str}",
                    key=f"dialog_chat_{idx}",
                    use_container_width=True,
                ):
                    st.session_state.selected_index = idx
                    st.session_state.current_page = "main"
                    st.rerun()
                st.caption(f"미리보기: {preview}")
                st.divider()


# ==========================================
# 5. 사이드바
# ==========================================
import streamlit as st

# 사이드바에 배치할 경우 st.file_uploader 대신 st.sidebar.file_uploader 사용
uploaded_file = st.file_uploader("별표 및 참고 문서 업로드 (PDF 전용)", type=["pdf"])

# 업로드된 파일 데이터를 세션 상태에 저장하여 processor.py에서 접근 가능하도록 설정
if uploaded_file is not None:
    st.session_state.uploaded_pdf_bytes = uploaded_file.getvalue()
else:
    st.session_state.uploaded_pdf_bytes = None
# 위의 코드는 파일 업로드 기능을 추가하기 위해 삽입되었음.
with st.sidebar:
    logo_path = find_yongin_logo()
    if logo_path:
        st.image(logo_path, width=172)
    st.markdown(
        "<div class='sidebar-brand-title'>용인시 건축법규 지원엔진</div>",
        unsafe_allow_html=True,
    )
    st.divider()

    with st.expander("실무자 로그인", expanded=(st.session_state.user_id == "guest")):
        if st.session_state.user_id == "guest":
            auth_tabs = st.tabs(["로그인", "계정 생성"])
            with auth_tabs[0]:
                login_id = st.text_input("아이디", key="login_id")
                login_pw = st.text_input("암호", type="password", key="login_pw")
                if st.button("로그인", use_container_width=True):
                    from storage import authenticate_user

                    if authenticate_user(login_id.strip(), login_pw.strip()):
                        st.session_state.user_id = login_id.strip()
                        st.session_state.chat_history = load_history(st.session_state.user_id)
                        st.session_state.selected_index = None
                        st.toast(f"{st.session_state.user_id} 계정으로 로그인했습니다.")
                        st.rerun()
                    else:
                        render_notice("아이디 또는 암호가 일치하지 않습니다.", "error")
            with auth_tabs[1]:
                reg_id = st.text_input("아이디", key="reg_id")
                reg_pw = st.text_input("암호", type="password", key="reg_pw")
                if st.button("계정 생성", use_container_width=True):
                    from storage import check_id_exists, register_user

                    cleaned_id = reg_id.strip()
                    cleaned_pw = reg_pw.strip()
                    if not cleaned_id or not cleaned_pw:
                        render_notice("아이디와 암호를 모두 입력해 주세요.", "warning")
                    elif check_id_exists(cleaned_id):
                        render_notice("이미 사용 중인 아이디입니다.", "error")
                    elif register_user(cleaned_id, cleaned_pw):
                        st.toast("계정이 생성되었습니다. 로그인해 주세요.")
        else:
            st.markdown(
                f"현재 사용자: <span style='color:var(--accent-strong); font-weight:700;'>{html.escape(st.session_state.user_id)}</span>",
                unsafe_allow_html=True,
            )
            if st.button("로그아웃", use_container_width=True, type="secondary"):
                st.session_state.user_id = "guest"
                st.session_state.chat_history = load_history("guest")
                st.session_state.selected_index = None
                st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)

    st.subheader("주요 메뉴")
    if st.button("AI 법규 검토", use_container_width=True):
        st.session_state.current_page = "main"
        st.rerun()
    if st.button("민원 서식 작성", use_container_width=True):
        st.session_state.current_page = "doc_gen"
        st.rerun()
    if st.button("실무 Q&A", use_container_width=True):
        st.session_state.current_page = "qna"
        st.rerun()
    if st.button("사이트맵", use_container_width=True):
        st.session_state.current_page = "sitemap"
        st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)
    dark_toggle_col, dark_label_col = st.columns([0.28, 0.72])
    with dark_toggle_col:
        st.toggle(
            "다크 모드",
            key="dark_mode_toggle",
            on_change=sync_dark_mode,
            label_visibility="collapsed",
        )
    with dark_label_col:
        st.markdown("<div class='toggle-caption'>다크 모드</div>", unsafe_allow_html=True)
    st.divider()

    st.subheader("최근 대화")
    c_new, c_src = st.columns(2)
    with c_new:
        if st.button("새 대화", use_container_width=True, type="primary"):
            st.session_state.selected_index = None
            st.session_state.current_page = "main"
            st.rerun()
    with c_src:
        if st.button("검색", use_container_width=True):
            open_history_search_dialog()

    history_container = st.container(height=400, border=True)
    with history_container:
        if st.session_state.chat_history:
            ensure_chat_ids()

            history_items = list(enumerate(st.session_state.chat_history))
            history_items = sorted(
                history_items,
                key=lambda item: (
                    item[1].get("pinned", False),
                    item[1].get("updated_at", item[1].get("created_at", "")),
                ),
                reverse=True,
            )

            for actual_index, chat in history_items:
                chat_id = chat["id"]
                time_str = chat.get("updated_at", chat.get("created_at", "00-00 00:00"))[5:16]
                title = chat.get("title", "질의")
                is_pinned = chat.get("pinned", False)
                pin_mark = "[고정] " if is_pinned else ""
                query_summary = title[:12] + ".." if len(title) > 12 else title

                col_btn, col_menu = st.columns([78, 22])
                with col_btn:
                    if st.button(
                        f"{pin_mark}{time_str} | {query_summary}",
                        key=f"hist_btn_{chat_id}",
                        use_container_width=True,
                    ):
                        st.session_state.selected_index = actual_index
                        st.session_state.current_page = "main"
                        st.rerun()
                with col_menu:
                    with st.popover(
                        "관리",
                        use_container_width=True,
                        help=f"대화 관리 {chat_id}",
                    ):
                        new_title = st.text_input(
                            "이름 변경",
                            value=title,
                            key=f"sidebar_rename_title_{chat_id}",
                        )

                        if st.button(
                            "이름 저장",
                            key=f"sidebar_save_title_{chat_id}",
                            use_container_width=True,
                        ):
                            st.session_state.chat_history[actual_index]["title"] = (
                                new_title.strip() or "질의"
                            )
                            save_history(st.session_state.chat_history, st.session_state.user_id)
                            st.toast("대화 이름을 변경했습니다.")
                            st.rerun()

                        pin_label = "고정 해제" if is_pinned else "상단 고정"

                        if st.button(
                            pin_label,
                            key=f"sidebar_toggle_pin_{chat_id}",
                            use_container_width=True,
                        ):
                            st.session_state.chat_history[actual_index]["pinned"] = not is_pinned
                            save_history(st.session_state.chat_history, st.session_state.user_id)
                            st.toast("대화 고정 상태를 변경했습니다.")
                            st.rerun()

                        if st.button(
                            "삭제",
                            key=f"sidebar_delete_chat_{chat_id}",
                            use_container_width=True,
                        ):
                            st.session_state.chat_history.pop(actual_index)

                            if st.session_state.selected_index == actual_index:
                                st.session_state.selected_index = None
                            elif (
                                st.session_state.selected_index is not None
                                and st.session_state.selected_index > actual_index
                            ):
                                st.session_state.selected_index -= 1

                            save_history(st.session_state.chat_history, st.session_state.user_id)
                            st.toast("대화를 삭제했습니다.")
                            st.rerun()
        else:
            st.caption("저장된 대화가 없습니다.")

    if st.session_state.chat_history:
        if st.button("전체 대화 삭제", type="secondary", use_container_width=True):
            st.session_state.chat_history = []
            st.session_state.selected_index = None
            from storage import clear_history

            clear_history(st.session_state.user_id)
            st.toast("저장된 대화를 모두 삭제했습니다.")
            st.rerun()


# ==========================================
# 6. 화면 1: AI 법규 검토
# ==========================================
if st.session_state.current_page == "main":
    st.markdown("## 건축 조례 및 규제 검토")
    st.caption("용인시 조례와 관련 법령을 함께 검토해 주요 기준, 예외 조건, 행정 검토 포인트를 정리합니다.")
    st.write("")

    if st.session_state.selected_index is not None:
        current_chat = st.session_state.chat_history[st.session_state.selected_index]
        if "state" not in current_chat:
            current_chat["state"] = {}
        current_state = current_chat["state"]
    else:
        current_chat = None
        current_state = {}

    col_top_left, col_top_right = st.columns([3.7, 1.3])
    with col_top_right:
        panel_toggle_col, panel_label_col = st.columns([0.25, 0.75])
        with panel_toggle_col:
            show_state_panel = st.toggle(
                "검토 조건 패널",
                value=True,
                key="use_state_panel",
                label_visibility="collapsed",
            )
        with panel_label_col:
            st.markdown("<div class='toggle-caption'>정보 관리창 On/Off</div>", unsafe_allow_html=True)

    if show_state_panel:
        col_chat, col_state = st.columns([73, 27])
    else:
        col_chat = st.container()
        col_state = None

    with col_chat:
        chat_box = st.container(height=560, border=True)

        user_query = st.chat_input(
            "검토할 조례, 대지 조건, 행정 쟁점을 입력하세요. 예: 처인구 자연녹지지역 내 건폐율 특례 적용 여부"
        )

        with chat_box:
            if current_chat is not None:
                render_notice(f"불러온 대화: {current_chat.get('title', '이전 대화')}", "info")
                for msg in current_chat.get("messages", []):
                    render_user_message(msg.get("query", ""))
                    render_ai_report(msg.get("response", ""))

                if st.button("현재 대화를 닫고 새 질의 시작", use_container_width=True):
                    st.session_state.selected_index = None
                    st.rerun()
            else:
                st.markdown(
                    """
                    <div class="app-intro">
                        <h3>검토할 건축 규제를 입력해 주세요</h3>
                        <p>상위 법령, 경기도 조례, 용인시 조례를 함께 검토해 실무 검토에 필요한 기준을 정리합니다.</p>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                c_info1, c_info2 = st.columns(2)
                with c_info1:
                    render_notice("질의 예시: 용인시 기흥구 상업지역 용적률 인센티브 완화 적용 범위", "info")
                with c_info2:
                    render_notice("질의 예시: 대지안의 공지 규정 관련 용인시 조례상 이격거리 예외 기준", "info")

        if user_query:
            render_user_message(user_query)
            with st.status("검토 중", expanded=True) as status:
                try:
                    st.write("관련 조례 및 법령을 확인하고 답변을 작성하는 중입니다.")
                    response_text = handle_ai_analysis(user_query)

                    if st.session_state.chat_history:
                        st.session_state.selected_index = len(st.session_state.chat_history) - 1

                    status.update(label="검토 완료", state="complete")
                    render_ai_report(response_text)
                    st.toast("AI 검토 결과가 생성되었습니다.")
                except Exception as e:
                    status.update(label="처리 중 오류가 발생했습니다", state="error")
                    render_notice(f"오류 내용: {str(e)}", "error")
            st.rerun()

    if show_state_panel and col_state is not None:
        with col_state:
            with st.container(border=True):
                st.markdown(
                    "<div class='context-panel-title'>토지 및 건축정보 관리창</div>",
                    unsafe_allow_html=True,
                )
                st.caption("아래 저장된 토지 및 건축정보는 대화에 지속적으로 반영됩니다.")

                

                if current_chat is not None:
                    if not current_state:
                        df_state = pd.DataFrame(columns=["변수명", "속성값"])
                    else:
                        df_state = pd.DataFrame(list(current_state.items()), columns=["변수명", "속성값"])

                    edited_df = st.data_editor(
                        df_state,
                        num_rows="dynamic",
                        use_container_width=True,
                        hide_index=True,
                        key=f"state_editor_v3_{st.session_state.selected_index}",
                    )

                    new_state = {}
                    for _, row in edited_df.iterrows():
                        k = str(row["변수명"]).strip() if pd.notna(row["변수명"]) else ""
                        v = str(row["속성값"]).strip() if pd.notna(row["속성값"]) else ""
                        if k and k != "nan":
                            new_state[k] = v

                    if new_state != current_state:
                        current_chat["state"] = new_state
                        save_history(st.session_state.chat_history, st.session_state.user_id)
                        st.rerun()
                else:
                    render_notice("대화가 선택되어 있지 않습니다. 새 질의를 작성하거나 왼쪽에서 대화를 선택해 주세요.", "info")


# ==========================================
# 7. 화면 2: 민원 서식 작성
# ==========================================
elif st.session_state.current_page == "doc_gen":
    st.markdown("## 민원 서식 작성")
    st.caption("작성하려는 민원 내용을 입력하면 행정 문서 형식에 맞게 정리합니다.")
    st.divider()

    required_docs = {
        "건축허가 관련": ["건축허가 신청서", "배치도", "평면도", "토지이용계획확인서", "건축계획서"],
        "건축선 문의": ["대지 위치도", "토지이용계획확인서", "현장 사진"],
        "일조권 민원": ["현장 사진", "건축물 배치도", "피해 설명 자료"],
        "불법건축물 신고": ["현장 사진", "위치도", "불법사항 설명자료"],
        "용도변경 문의": ["건축물대장", "평면도", "용도변경 계획서"],
        "주차장 기준 문의": ["배치도", "주차계획도", "건축 개요"],
        "건축물 해석 문의": ["질의서", "관련 도면", "현장 사진"],
        "기타": ["신분증", "민원 설명자료"],
    }
    department_map = {
        "건축허가 관련": "건축허가과",
        "건축선 문의": "건축과",
        "일조권 민원": "건축과",
        "불법건축물 신고": "건축과",
        "용도변경 문의": "건축허가과",
        "주차장 기준 문의": "교통정책과",
        "건축물 해석 문의": "건축과",
        "기타": "민원여권과",
    }

    col1, col2 = st.columns(2)
    with col1:
        civil_type = st.selectbox("처리 유형", list(required_docs.keys()))
    with col2:
        site_address = st.text_input("대상지 주소", placeholder="예: 경기도 용인시 처인구 역북동")

    civil_content = st.text_area(
        "요청 내용",
        height=140,
        placeholder="민원으로 정리할 사실관계, 요청 사항, 확인이 필요한 쟁점을 입력해 주세요.",
    )

    if st.button("서식 작성", use_container_width=True, type="primary"):
        if not site_address or not civil_content:
            render_notice("대상지 주소와 요청 내용을 모두 입력해 주세요.", "warning")
        else:
            with st.spinner("민원 서식을 작성하는 중입니다."):
                try:
                    result = generate_civil_document(civil_type, site_address, civil_content)
                    st.session_state.selected_index = None
                    st.toast("민원 서식이 작성되었습니다.")

                    st.divider()
                    st.subheader("작성된 민원 서식")
                    with st.container(border=True):
                        st.markdown(result)

                    info_tabs = st.tabs(["필수 서류", "담당 부서", "접수 절차", "파일 다운로드"])

                    with info_tabs[0]:
                        for doc_name in required_docs.get(civil_type, required_docs["기타"]):
                            st.write(f"- {doc_name}")
                    with info_tabs[1]:
                        dept = department_map.get(civil_type, "민원여권과")
                        render_notice(f"용인시청 또는 관할 구청 담당 부서: {dept}", "info")
                    with info_tabs[2]:
                        st.markdown(
                            """
                            1. 온라인 접수: 정부24 또는 세움터를 통한 전자민원 접수
                            2. 방문 접수: 용인시청 또는 관할 구청 민원 창구 방문 접수
                            3. 사전 확인: 도면, 사진, 토지 관련 서류 등 첨부자료 누락 여부 확인
                            """
                        )
                    with info_tabs[3]:
                        doc = Document()
                        doc.add_heading("용인시 건축 행정 민원 문서", level=1)
                        doc.add_paragraph(f"처리 유형: {civil_type}\n대상지 주소: {site_address}")
                        doc.add_heading("본문", level=2)
                        doc.add_paragraph(result)

                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        st.download_button(
                            label="DOCX 다운로드",
                            data=buffer,
                            file_name=f"용인시_건축서식_{civil_type}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                except Exception as e:
                    render_notice(f"문서 생성 중 오류가 발생했습니다: {str(e)}", "error")


# ==========================================
# 8. 화면 3: 실무 Q&A
# ==========================================
elif st.session_state.current_page == "qna":
    st.markdown("## 프로그램 Q&A")
    st.caption("추가 기능 또는 법규 목록 및 기타 질의사항을 남겨주시면 답변드리겠습니다.")
    st.divider()

    col1, col2 = st.columns([65, 35])
    with col1:
        st.subheader("질의응답 목록")
        if not st.session_state.qna_list:
            render_notice("등록된 질문이 없습니다.", "info")
        else:
            for i, q in enumerate(st.session_state.qna_list):
                badge = "대기" if q["status"] == "대기중" else "답변 완료"
                with st.expander(f"[{badge}] {q['title']}"):
                    st.markdown(f"**질의 요지:** {q['content']}")
                    if q["status"] == "답변완료":
                        render_notice(f"답변: {q['answer']}", "success")

    with col2:
        st.subheader("질문 등록")
        with st.form("qna_form_v3", clear_on_submit=True):
            q_title = st.text_input("제목")
            q_content = st.text_area("내용")
            if st.form_submit_button("질문 등록", use_container_width=True):
                if q_title and q_content:
                    st.session_state.qna_list.append(
                        {
                            "title": q_title,
                            "content": q_content,
                            "status": "대기중",
                            "answer": "",
                        }
                    )
                    st.toast("질문이 등록되었습니다.")
                    st.rerun()
                else:
                    render_notice("제목과 내용을 모두 입력해 주세요.", "warning")

        st.divider()
        with st.expander("관리자 답변"):
            admin_pw = st.text_input("관리자 PIN", type="password")
            if admin_pw == "2026":
                render_notice("관리자 권한이 확인되었습니다.", "success")
                for i, q in enumerate(st.session_state.qna_list):
                    if q["status"] == "대기중":
                        answer_text = st.text_input(f"답변 입력: {q['title']}", key=f"ans_{i}")
                        if st.button("답변 등록", key=f"btn_{i}"):
                            st.session_state.qna_list[i]["answer"] = answer_text
                            st.session_state.qna_list[i]["status"] = "답변완료"
                            st.rerun()


# ==========================================
# 9. 화면 4: 사이트맵
# ==========================================
elif st.session_state.current_page == "sitemap":
    st.markdown("## 사이트맵")
    st.caption("플랫폼의 주요 기능, 분석 엔진, 학습 법규 데이터 범위를 한눈에 확인할 수 있습니다.")
    st.divider()

    architecture_html = textwrap.dedent(
        """
        <!DOCTYPE html>
        <html lang="ko">
        <head>
        <meta charset="UTF-8">
        <style>
            @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');

            :root {
                --bg: #f5f5f5;
                --surface: #ffffff;
                --surface-2: #fafafa;
                --text: #111111;
                --muted: #555555;
                --line: #b8b8b8;
                --accent: #111111;
                --green: #333333;
                --violet: #222222;
            }

            * {
                box-sizing: border-box;
            }

            body {
                margin: 0;
                padding: 18px;
                background: var(--bg);
                font-family: Pretendard, -apple-system, BlinkMacSystemFont, sans-serif;
                color: var(--text);
            }

            .wrapper {
                background: var(--surface);
                border: 1.25px solid var(--line);
                border-radius: 18px;
                padding: 26px;
                box-shadow: 0 12px 30px rgba(17, 17, 17, 0.06);
            }

            .title {
                text-align: center;
                font-size: 26px;
                font-weight: 800;
                margin-bottom: 8px;
            }

            .subtitle {
                text-align: center;
                color: var(--muted);
                font-size: 14px;
                margin-bottom: 26px;
            }

            .section {
                background: var(--surface-2);
                border: 1.25px solid var(--line);
                border-radius: 14px;
                padding: 18px;
                margin-bottom: 16px;
            }

            .section-title {
                font-weight: 800;
                color: var(--text);
                font-size: 17px;
                margin-bottom: 14px;
            }

            .grid {
                display: grid;
                grid-template-columns: repeat(4, minmax(0, 1fr));
                gap: 12px;
            }

            .card {
                background: var(--surface);
                border: 1.25px solid var(--line);
                border-radius: 12px;
                padding: 14px;
                min-height: 96px;
            }

            .card-title {
                font-weight: 800;
                font-size: 15px;
                margin-bottom: 8px;
            }

            .blue .card-title {
                color: var(--accent);
            }

            .green .card-title {
                color: var(--green);
            }

            .violet .card-title {
                color: var(--violet);
            }

            .card-desc {
                font-size: 13px;
                color: var(--muted);
                line-height: 1.55;
            }

            @media (max-width: 900px) {
                .grid {
                    grid-template-columns: repeat(2, minmax(0, 1fr));
                }
            }

            @media (max-width: 540px) {
                .grid {
                    grid-template-columns: 1fr;
                }
            }
        </style>
        </head>

        <body>
            <div class="wrapper">
                <div class="title">용인시 건축 조례 지원 플랫폼 구조</div>
                <div class="subtitle">실무 검토 화면, 분석 엔진, 법규 데이터의 연결 구조</div>

                <div class="section">
                    <div class="section-title">사용자 업무 화면</div>
                    <div class="grid">
                        <div class="card blue">
                            <div class="card-title">AI 법규 검토</div>
                            <div class="card-desc">조례와 상위 법령을 함께 검토하는 질의응답 화면</div>
                        </div>
                        <div class="card blue">
                            <div class="card-title">민원 서식 작성</div>
                            <div class="card-desc">입력 내용을 행정 문서 형식으로 정리하는 작성 도구</div>
                        </div>
                        <div class="card blue">
                            <div class="card-title">실무 Q&A</div>
                            <div class="card-desc">업무 중 발생한 쟁점과 답변을 축적하는 공유 공간</div>
                        </div>
                        <div class="card blue">
                            <div class="card-title">검토 조건 관리</div>
                            <div class="card-desc">대지 조건과 행정 변수를 대화별로 관리하는 보조 패널</div>
                        </div>
                    </div>
                </div>

                <div class="section">
                    <div class="section-title">AI 및 백엔드 엔진</div>
                    <div class="grid">
                        <div class="card violet">
                            <div class="card-title">분석 엔진</div>
                            <div class="card-desc">handle_ai_analysis 모듈을 통한 질의 분석과 답변 생성</div>
                        </div>
                        <div class="card violet">
                            <div class="card-title">법규 매핑</div>
                            <div class="card-desc">조례, 시행령, 시행규칙의 인용 관계를 기준으로 검토 범위 구성</div>
                        </div>
                        <div class="card violet">
                            <div class="card-title">세션 관리</div>
                            <div class="card-desc">사용자별 대화 이력, 선택 대화, 검토 조건 저장</div>
                        </div>
                        <div class="card violet">
                            <div class="card-title">문서 생성</div>
                            <div class="card-desc">민원 내용을 표준 문서 형식과 DOCX 파일로 변환</div>
                        </div>
                    </div>
                </div>

                <div class="section">
                    <div class="section-title">학습 및 참조 데이터</div>
                    <div class="grid">
                        <div class="card green">
                            <div class="card-title">용인시 조례 및 경기도 조례</div>
                            <div class="card-desc">건축 조례, 건축물관리 조례, 도시계획 조례 및 경기도 건축 조례 등</div>
                        </div>
                        <div class="card green">
                            <div class="card-title">용인시 홈페이지 사이트맵</div>
                            <div class="card-desc">질문에 관련된 용인시 홈페이지가 있을 경우 안내</div>
                        </div>
                        <div class="card green">
                            <div class="card-title">상위 법령</div>
                            <div class="card-desc">건축법, 건축기본법, 국토계획법 및 하위 법령</div>
                        </div>
                        <div class="card green">
                            <div class="card-title">차용 법규</div>
                            <div class="card-desc">조례와 법령에서 준용하거나 인용하는 관련 법규</div>
                        </div>
                    </div>
                </div>
            </div>
        </body>
        </html>
        """
    )

    st_components.html(
        architecture_html,
        height=760,
        scrolling=True,
    )

    st.divider()
    st.subheader("학습 법규 목록")
    st.caption("플랫폼 검토 범위에 포함된 기본 조례, 상위 법령, 차용 법규를 정리했습니다.")

    tabs = st.tabs(["기본 조례", "상위 기본 법령", "조례의 차용 법규", "법령의 차용 법규"])

    with tabs[0]:
        df_ord = pd.DataFrame(
            [
                ["경기도", "경기도 건축 조례"],
                ["경기도", "경기도 건축기본조례"],
                ["경기도", "경기도 건축물 미술작품 설치 및 관리에 관한 조례"],
                ["경기도", "경기도 건축물관리 조례"],
                ["용인시", "용인시 건축 조례"],
                ["용인시", "용인시 건축물관리 조례"],
                ["용인시", "용인시 도시계획 조례"],
            ],
            columns=["소관", "조례명"],
        )
        render_legal_table(df_ord)

    with tabs[1]:
        law_data = [
            "건축법",
            "건축기본법",
            "문화예술진흥법",
            "건축물관리법",
            "국토의 계획 및 이용에 관한 법률",
            "건축법 시행령",
            "건축기본법 시행령",
            "문화예술진흥법 시행령",
            "건축물관리법 시행령",
            "국토의 계획 및 이용에 관한 법률 시행령",
            "건축법 시행규칙",
            "건축물관리법 시행규칙",
            "국토의 계획 및 이용에 관한 법률 시행규칙",
        ]
        render_legal_table(to_multi_column_dataframe(law_data, column_count=3, column_name="법령"))

    with tabs[2]:
        ord_borrow_data = [
            "감정평가 및 감정평가사에 관한 법률",
            "건설기술진흥법",
            "건축기본법",
            "건축물관리법",
            "건축법",
            "건축사법",
            "고등교육법",
            "공공주택 특별법",
            "관광진흥법",
            "국가유산기본법",
            "국토의 계획 및 이용에 관한 법률",
            "근현대문화유산의 보존 및 활용에 관한 법률",
            "녹색건축물 조성 지원법",
            "농지법",
            "대기환경보전법",
            "도시 및 주거환경정비법",
            "문화예술진흥법",
            "민간임대주택에 관한 특별법",
            "산업입지 및 개발에 관한 법률",
            "산업집적활성화 및 공장설립에 관한 법률",
            "산지관리법",
            "실내공기질 관리법",
            "자연공원법",
            "전통사찰의 보존 및 지원에 관한 법률",
            "주택법",
            "한옥 등 건축자산의 진흥에 관한 법률",
            "건축기본법 시행령",
            "건축물관리법 시행령",
            "건축법 시행령",
            "국토의 계획 및 이용에 관한 법률 시행령",
            "농지법 시행령",
            "다중이용업소의 안전관리에 관한 특별법 시행령",
            "문화예술진흥법 시행령",
            "민원 처리에 관한 법률 시행령",
            "건축물관리법 시행규칙",
            "건축법 시행규칙",
            "국토의 계획 및 이용에 관한 법률 시행규칙",
            "경기도 건축물 미술작품 설치 및 관리에 관한 조례 시행규칙",
            "경기도 위원회 수당 및 여비 지급 조례",
            "경기도 지방보조금 관리 조례",
            "용인시 각종 위원회 설치 및 운영 조례",
            "용인시 건축 조례",
            "용인시 도시계획 조례",
            "공공발주사업에 대한 건축사의 업무범위와 대가기준",
        ]
        render_legal_table(to_multi_column_dataframe(ord_borrow_data, column_count=4, column_name="법규"))

    with tabs[3]:
        law_borrow_data = [
            "건설기술 진흥법",
            "건설산업기본법",
            "건축기본법",
            "건축법",
            "건축사법",
            "경관법",
            "고등교육법",
            "공간정보의 구축 및 관리 등에 관한 법률",
            "공공기관의 운영에 관한 법률",
            "공공주택 특별법",
            "공동주택관리법",
            "공유재산 및 물품 관리법",
            "공익사업을 위한 토지 등의 취득 및 보상에 관한 법률",
            "관광진흥법",
            "국가기술자격법",
            "국가유산기본법",
            "국유재산법",
            "국토안전관리원법",
            "국토의 계획 및 이용에 관한 법률",
            "기술사법",
            "녹색건축물 조성 지원법",
            "농어촌정비법",
            "농지법",
            "도로법",
            "도시 및 주거환경정비법",
            "도시개발법",
            "도시공원 및 녹지 등에 관한 법률",
            "도시교통정비 촉진법",
            "도시재생 활성화 및 지원에 관한 특별법",
            "문화예술진흥법",
            "문화유산의 보존 및 활용에 관한 법률",
            "민법",
            "빈집 및 소규모주택 정비에 관한 특례법",
            "사도법",
            "산림자원의 조성 및 관리에 관한 법률",
            "산업입지 및 개발에 관한 법률",
            "산업집적활성화 및 공장설립에 관한 법률",
            "산지관리법",
            "소방시설 설치 및 관리에 관한 법률",
            "수도권정비계획법",
            "수도법",
            "시설물의 안전 및 유지관리에 관한 특별법",
            "영유아보육법",
            "자연공원법",
            "자연유산의 보존 및 활용에 관한 법률",
            "자연재해대책법",
            "전자정부법",
            "주차장법",
            "주택법",
            "지방공기업법",
            "집합건물의 소유 및 관리에 관한 법률",
            "택지개발촉진법",
            "토지이용규제 기본법",
            "하수도법",
            "하천법",
            "한국토지주택공사법",
            "행정대집행법",
            "건설기술 진흥법 시행령",
            "건축법 시행령",
            "국토의 계획 및 이용에 관한 법률 시행령",
            "수산자원관리법 시행령",
        ]
        render_legal_table(to_multi_column_dataframe(law_borrow_data, column_count=4, column_name="법규"))

import requests
import urllib.parse
import time
import streamlit as st

def check_national_law_updates(api_key, base_date=20260403):
    national_laws = [
        "건축법", "건축기본법", "문화예술진흥법", "건축물관리법", "국토의 계획 및 이용에 관한 법률",
        "감정평가 및 감정평가사에 관한 법률", "건설기술 진흥법", "건축사법", "고등교육법", 
        "공공주택 특별법", "관광진흥법", "국가유산기본법", "근현대문화유산의 보존 및 활용에 관한 법률", 
        "녹색건축물 조성 지원법", "농지법", "대기환경보전법", "도시 및 주거환경정비법", 
        "민간임대주택에 관한 특별법", "산업입지 및 개발에 관한 법률", "산업집적활성화 및 공장설립에 관한 법률", 
        "산지관리법", "실내공기질 관리법", "자연공원법", "전통사찰의 보존 및 지원에 관한 법률", "주택법",
        "한옥 등 건축자산의 진흥에 관한 법률", "민원 처리에 관한 법률", "건설산업기본법", "경관법", 
        "공간정보의 구축 및 관리 등에 관한 법률", "공공기관의 운영에 관한 법률", "공동주택관리법", 
        "공유재산 및 물품 관리법", "공익사업을 위한 토지 등의 취득 및 보상에 관한 법률", 
        "국가기술자격법", "국유재산법", "국토안전관리원법", "기술사법", "농어촌정비법", "도로법", 
        "도시개발법", "도시공원 및 녹지 등에 관한 법률", "도시교통정비 촉진법",
        "도시재생 활성화 및 지원에 관한 특별법", "문화유산의 보존 및 활용에 관한 법률", "민법",
        "빈집 및 소규모주택 정비에 관한 특례법", "사도법", "산림자원의 조성 및 관리에 관한 법률",
        "소방시설 설치 및 관리에 관한 법률", "수도권정비계획법", "수도법",
        "시설물의 안전 및 유지관리에 관한 특별법", "영유아보육법",
        "자연유산의 보존 및 활용에 관한 법률", "자연재해대책법", "전자정부법", "주차장법",
        "지방공기업법", "집합건물의 소유 및 관리에 관한 법률", "택지개발촉진법",
        "토지이용규제 기본법", "하수도법", "하천법", "한국토지주택공사법", "행정대집행법",
        "공공발주사업에 대한 건축사의 업무범위와 대가기준"
    ]

    updated_laws = []
    slash = chr(47)
    agent_string = f"Mozilla_5.0 (Windows NT 10.0; Win64; x64) AppleWebKit_537.36 (KHTML, like Gecko) Chrome_120.0.0.0 Safari_537.36".replace("_", slash)
    headers = {"User-Agent": agent_string}

    chunk_size = 10
    for i in range(0, len(national_laws), chunk_size):
        chunk = national_laws[i:i + chunk_size]
        for law in chunk:
            encoded_law = urllib.parse.quote(law)
            url = f"http:{slash}{slash}www.law.go.kr{slash}DRF{slash}lawSearch.do?OC={api_key}&target=oneview&type=JSON&query={encoded_law}"
            
            try:
                res = requests.get(url, headers=headers, timeout=10)
                if res.status_code == 200:
                    data = res.json()
                    items_block = data.get("items", {})
                    
                    total_cnt = items_block.get("totalCnt", "0")
                    if total_cnt != "0" and "item" in items_block:
                        for law_item in items_block["item"]:
                            law_name = law_item.get("법령명", "")
                            efyd = law_item.get("시행일자", "")
                            
                            if efyd and int(efyd) > base_date:
                                if law_name not in updated_laws:
                                    updated_laws.append(law_name)
            except Exception:
                continue
                
            time.sleep(1.5)
            
    return updated_laws

def check_local_ordinance_updates(api_key, base_date=20260425):
    import requests
    import urllib.parse
    import time

    local_ordinances = [
        "경기도 건축 조례", "경기도 건축기본조례", "경기도 건축물 미술작품 설치 및 관리에 관한 조례",
        "경기도 건축물관리 조례", "용인시 건축 조례", "용인시 건축물관리 조례",
        "용인시 도시계획 조례", "경기도 건축물 미술작품 설치 및 관리에 관한 조례 시행규칙",
        "경기도 위원회 수당 및 여비 지급 조례", "경기도 지방보조금 관리 조례",
        "용인시 각종 위원회 설치 및 운영 조례"
    ]

    updated_ordinances = []
    slash = chr(47)
    agent_string = f"Mozilla_5.0 (Windows NT 10.0; Win64; x64) AppleWebKit_537.36 (KHTML, like Gecko) Chrome_120.0.0.0 Safari_537.36".replace("_", slash)
    headers = {"User-Agent": agent_string}

    chunk_size = 10
    for i in range(0, len(local_ordinances), chunk_size):
        chunk = local_ordinances[i:i + chunk_size]
        for ordin in chunk:
            search_keyword = ordin.replace(" ", "")
            encoded_ordin = urllib.parse.quote(search_keyword)
            
            url = f"http:{slash}{slash}www.law.go.kr{slash}DRF{slash}lawSearch.do?OC={api_key}&target=ordin&type=JSON&query={encoded_ordin}"
            
            try:
                res = requests.get(url, headers=headers, timeout=10)
                if res.status_code == 200:
                    data = res.json()
                    
                    # 1. 맹점 해결: 폴더 이름이 무엇이든 'totalCnt'가 있는 핵심 폴더를 동적으로 추적
                    root_block = data
                    for key, value in data.items():
                        if isinstance(value, dict) and "totalCnt" in value:
                            root_block = value
                            break
                            
                    total_cnt = root_block.get("totalCnt", "0")
                    
                    if str(total_cnt) != "0":
                        item_list = []
                        if "ordin" in root_block:
                            item_list = root_block["ordin"]
                        elif "item" in root_block:
                            item_list = root_block["item"]
                            
                        # 2. 맹점 해결: 반환된 모든 버전을 순회하며 기준일자 이후의 개정안이 있는지 스캔
                        if item_list:
                            for ordin_item in item_list:
                                efyd = ordin_item.get("시행일자", ordin_item.get("efYd", ""))
                                
                                # 데이터에 섞여 있을지 모를 공백이나 특수문자를 제거하고 순수 숫자만 추출
                                efyd_clean = "".join(filter(str.isdigit, str(efyd)))
                                
                                if efyd_clean and int(efyd_clean) > base_date:
                                    if ordin not in updated_ordinances:
                                        updated_ordinances.append(ordin)
                                    break  # 개정 사실을 확인했으므로 즉시 중단하고 다음 조례로 넘어감
            except Exception:
                continue
                
            time.sleep(1.5)
            
    return updated_ordinances

# 배포 시 아래 두 함수의 주석 샵 기호를 제거한다.
# @st.cache_data(ttl=86400) 
def run_daily_national_update_check(date_val):
    api_key = st.secrets["LAW_API_KEY"]
    return check_national_law_updates(api_key, base_date=date_val)

# 배포 시 아래 두 함수의 주석 샵 기호를 제거한다.
# @st.cache_data(ttl=86400) 
def run_daily_ordinance_update_check(date_val):
    api_key = st.secrets["LAW_API_KEY"]
    return check_local_ordinance_updates(api_key, base_date=date_val)


test_base_date = 20140101

st.write("---")

col_main, col_alert = st.columns([7, 3])

with col_alert:
    st.subheader("법규 개정 모니터링(2014년 1월 1일 기준)")
    
    with st.spinner("2014년 1월 1일을 기준으로 법규 개정 여부를 확인하고 있습니다. 약 2분이 소요됩니다."):
        updated_list = run_daily_national_update_check(test_base_date)
        updated_ordinances_list = run_daily_ordinance_update_check(test_base_date)

    if updated_list:
        st.error("최신 개정안 시행 국가법령")
        for law_name in updated_list:
            st.write(f"- {law_name}")
    else:
        st.success("새롭게 개정된 국가법령이 없다.")

    st.write(" ")

    if updated_ordinances_list:
        st.error("최신 개정안 시행 자치법규")
        for ordin_name in updated_ordinances_list:
            st.write(f"- {ordin_name}")
    else:
        st.success("새롭게 개정된 자치법규가 없다.")

with col_main:
    st.write("메인 웹사이트 구역이다. 이 곳에 기존 메인 화면 구성 요소들이 위치해야 한다.")
